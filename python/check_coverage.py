#!/usr/bin/env python
'''
Copyright (C) 2015 The Regents of the University of California.

This program is free software: you can redistribute it and/or modify
it under the terms of the GNU General Public License as published by
the Free Software Foundation, either version 3 of the License, or
(at your option) any later version.

This program is distributed in the hope that it will be useful,
but WITHOUT ANY WARRANTY; without even the implied warranty of
MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
GNU General Public License for more details.

You should have received a copy of the GNU General Public License
along with this program.  If not, see <http://www.gnu.org/licenses/>.
'''

from __future__ import print_function

__author__ = "Alex Huszagh"
__maintainer__ = "Alex Huszagh"
__email__ = "ahuszagh@gmail.com"

# KNOWN ISSUES:
#   With enough proteins entered, the UniProt KB database will send
#   a 503 gateway error due to refused service.

# This program aims to analyze sequence coverage from a batch list
# of protein identifiers, with optional custom labels given.
# This is useful for refining experiments based on various digestive
# conditions, such as different proteases, different denaturants, etc.
# in order to optimize for a target bait sequence.

# Ex.:
# INPUT:
#   $ python check_coverage.py -c Trypsin Chymo -p P46406 P00761 -f a.txt b.txt
# This script also now runs a graphical user interface and exports as a
# a DOCX, or directly to the console.


# OUTPUT:
# -------------------------
# >sp|P46406|G3P_RABIT Glyceraldehyde-3-phosphate dehydrogenase OS=Oryctolagus cuniculus GN=GAPDH PE=1 SV=3
#
#             0: MVKVGVNGFGRIGRLVTRAAFNSGKVDVVAINDPFIDLHYMVYMFQYDSTHGKFHGTVKA
# Trypsin :      +++++++++++++++++++++++++                            +++++++
# Chymo :                                                             +++++++
#
#            60: ENGKLVINGKAITIFQERDPANIKWGDAGAEYVVESTGVFTTMEKAGAHLKGGAKRVIIS
# Trypsin :      ++++++++++++++++++++++++                     +++++++++++++++
# Chymo :        ++++
#
#           120: APSADAPMFVMGVNHEKYDNSLKIVSNASCTTNCLAPLAKVIHDHFGIVEGLMTTVHAIT
# Trypsin :      +++++++++++++++++
# Chymo :
#
#           180: ATQKTVDGPSGKLWRDGRGAAQNIIPASTGAAKAVGKVIPELNGKLTGMAFRVPTPNVSV
# Trypsin :          +++++++++++                              +++++++++++++++
# Chymo :
#
#           240: VDLTCRLEKAAKYDDIKKVVKQASEGPLKGILGYTEDQVVSCDFNSATHSSTFDAGAGIA
# Trypsin :      +++++++++++++++++++++++++++++
# Chymo :              +++++++++++ +++++++++++
#
#           300: LNDHFVKLISWYDNEFGYSNRVVDLMVHMASKE
# Trypsin :                           ++++++++++++
# Chymo :
#
# -------------------------
# >sp|P00761|TRYP_PIG Trypsin OS=Sus scrofa PE=1 SV=1
#
#             0: FPTDDDDKIVGGYTCAANSIPYQVSLNSGSHFCGGSLINSQWVVSAAHCYKSRIQVRLGE
# Trypsin :
# Chymo :
#
#            60: HNIDVLEGNEQFINAAKIITHPNFNGNTLDNDIMLIKLSSPATLNSRVATVSLPRSCAAA
# Trypsin :
# Chymo :
#
#           120: GTECLISGWGNTKSSGSSYPSLLQCLKAPVLSDSSCKSSYPGQITGNMICVGFLEGGKDS
# Trypsin :
# Chymo :
#
#           180: CQGDSGGPVVCNGQLQGIVSWGYGCAQKNKPGVYTKVCNYVNWIQQTIAAN
# Trypsin :
# Chymo :
#
# -------------------------

# load modules
import argparse
import os
import six
if six.PY2:
    import httplib
else:
    import http.client
import sys

import numpy as np
import pandas as pd
from PySide import QtCore, QtGui

# load objects/functions
from collections import namedtuple
from functools import partial

# CONSTANTS

QLABEL_BANNER_STYLE = '''
QLabel {
    font-weight: bold;
    font-size: 20pt;
};
'''

QLABEL_STYLE = '''
QLabel {
    font-weight: bold;
};
'''

SERVER = {
    'host': 'www.uniprot.org',
    'path': '/',
    'domain': 'uniprot',
    'suffix': '.fasta'
}
PATH = os.getcwd()
PROTEIN_COLUMN = 'Acc #'
PEPTIDE_COLUMN = 'DB Peptide'
ROW_HEIGHT = 50

# args
PARSER = argparse.ArgumentParser()
PARSER.add_argument("-c", "--conditions", nargs='+',
                    help="labels for experiments")
PARSER.add_argument("-f", "--files", nargs='+',
                    help="Input Tab-Delimited Text Files")
PARSER.add_argument("-p", "--protein", type=str, nargs='+',
                    help="UniProt ID")
PARSER.add_argument("-m", "--mode", type=str, default="docx",
                    choices=["text", "console", "docx"],
                    help="Mode, can print to console, write plain text"
                    " or write to a Open Document standard")
PARSER.add_argument("-o", "--output", type=str, default="out",
                    help="Name of output file")
ARGS = PARSER.parse_args()

# parse arguments
if ARGS.mode == 'docx':
    # need to load our modules
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.style import WD_STYLE_TYPE
    except ImportError:
        raise ImportError("check_coverage in docx mode requires"
                          "a python-docx installation, which can"
                          "be installed via pip.")

# ------------------
#       UTILS
# ------------------


def find_all(value, sub):
    """Find all occurrences within a string"""

    start = 0
    while True:
        start = value.find(sub, start)
        if start == -1:
            return
        yield start
        start += len(sub)


def uniquer(seq, idfun=None):
    '''
    Converts a sequence to a unique list while keeping order.
    Recipe modified from:
    https://code.activestate.com/recipes/52560-remove-duplicates-from-a-sequence/
    :
        >>> uniquer(range(4) + range(-2, 4))
        [0, 1, 2, 3, -2, -1]
    '''

    if idfun is None:
        # pylint: disable=E0102
        def idfun(var):
            return var
    seen = {}
    result = []
    for item in seq:
        marker = idfun(item)
        # in old Python versions:
        # if seen.has_key(marker)
        # but in new ones:
        if marker in seen:
            continue
        seen[marker] = 1
        result.append(item)
    return result


def block_once(widget, func):
    '''
    Blocks Qt signals for a single instance.
    :
        widget -- instance which inherits from QWidget
        func -- frozen function to call
    '''

    signal_state = widget.signalsBlocked()
    widget.blockSignals(True)
    func()
    widget.blockSignals(signal_state)

# ------------------
# PEPTIDE FUNCTIONS
# ------------------


def protein_coverage(dataframe, protein, sequence):
    '''Returns the protein coverage for a given UniProt ID bait and
    sequence.
    '''

    # init return
    sequence = ''.join(sequence[1:])
    data = {k: False for k in range(len(sequence))}
    cuts = {k: False for k in range(len(sequence))}
    # grab sequences
    indexes, = np.where(dataframe[PROTEIN_COLUMN] == protein)
    sequences = set(dataframe[PEPTIDE_COLUMN][indexes].tolist())
    # iterate over sequences
    for seq in sequences:
        positions = list(find_all(sequence, seq))
        for position in positions:
            for index in range(len(seq)):
                key = index + position
                data[key] = True
            # now need to add for the last one
            # used loop
            cuts[key] = True
    return data, cuts


def get_coverage(dataframes, protein, sequence):
    '''Iterativelt returns the protein coverage for each df within
    dataframes.
    '''

    # init return
    coverage_list = []
    cut_list = []
    for dataframe in dataframes:
        coverage, cuts = protein_coverage(dataframe, protein, sequence)
        coverage_list.append(coverage)
        cut_list.append(cuts)
    return coverage_list, cut_list


# ------------------
#    OUT FUNCTIONS
# ------------------


class Writer(object):
    '''Custom implementation of a console/text/docx Writer'''

    paragraph = None

    def __init__(self, mode, out):
        super(Writer, self).__init__()

        self.mode = mode
        if self.mode == 'text':
            path = self._get_path(out)
            if os.path.exists(os.path.dirname(path)):
                self.file = open(path, 'w')
        elif self.mode == 'console':
            self.file = sys.stdout
        elif self.mode == 'docx':
            path = self._get_path(out)
            if os.path.exists(os.path.dirname(path)):
                self.path = path
                self.file = Document()
                self._add_styles()

    # ------------------
    #        MAIN
    # ------------------

    def start_sequence(self, sequence):
        '''Starts the lines for a new protein'''

        if self.mode in ['text', 'console']:
            print('-------------------------', file=self.file)
            print(sequence[0], file=self.file)
            print(file=self.file)
        elif self.mode == 'docx':
            self.file.add_heading('-------------------------\n', 1)
            self.paragraph = self.file.add_paragraph(sequence[0] + '\n')
            self.paragraph.style = self.file.styles['Normal']

    def write_sequence_line(self, sequence, index, indent=15):
        '''Writes a sequence line with indentation to file'''

        # grab offset to add to file
        offset = str(len(''.join(sequence[1:index]))) + ': '
        offset_length = len(offset)
        # init line
        line = ' '*(indent-offset_length)
        line += offset
        line += sequence[index]
        if self.mode in ['text', 'console']:
            print(line, file=self.file)
        elif self.mode == 'docx':
            self.paragraph.add_run(line + '\n')

    def write_blank_line(self):
        '''Write blank line to file'''

        if self.mode in ['text', 'console']:
            print(file=self.file)
        elif self.mode == 'docx':
            self.paragraph.add_run('\n')

    def write_condition(self, header, coverage, cuts, sequence, index):
        '''Writes the condition with coverage to file'''

        output = self.process_condition(header, coverage, cuts,
                                        sequence, index)
        if self.mode in ['text', 'console']:
            print(output, file=self.file)
        elif self.mode == 'docx':
            self.paragraph.add_run('\n')

    def close_sequence(self):
        '''Closes the lines for a protein'''

        if self.mode in ['text', 'console']:
            print('-------------------------', file=self.file)
            print(file=self.file)
        elif self.mode == 'docx':
            self.file.add_heading('-------------------------\n', 1)

    # ------------------
    #       UTILS
    # ------------------

    @staticmethod
    def get_header(condition, total=12):
        '''Grabs a 35 character header from the given condition'''

        length = min([len(condition), total])
        header = condition[:total]
        header = ''.join([header, ' : '])
        header = ''.join([header, ' '*(total-length)])
        return header

    def close(self):
        '''Closes the writeable object'''

        if self.mode == 'text' and hasattr(self, "file"):
            self.file.close()
        elif self.mode == 'docx' and hasattr(self, "file"):
            self.file.save(self.path)

    def process_condition(self, header, coverage, cuts, sequence, index):
        '''
        Processes the header to give the conditions coverage of the
        sequence.
        '''

        # grab parameter lengths to determine range
        length = len(sequence[index])
        offset = len(''.join(sequence[1:index]))
        # iteratively add null string or +
        if self.mode == 'docx':
            self.paragraph.add_run(header)
        keys = range(offset, offset+length)
        for key in keys:
            value = coverage[key]
            cut = cuts[key]
            # add 'o' if cutsite, '+' if not, ' ' if blank
            if value and cut and self.mode in ['text', 'console']:
                header += 'o'
            elif value and not cut and self.mode in ['text', 'console']:
                header += '+'
            elif self.mode in ['text', 'console']:
                header += ' '
            # docx settings
            elif value and cut and self.mode == 'docx':
                self.paragraph.add_run('o', style='Red')
            elif value and not cut and self.mode == 'docx':
                self.paragraph.add_run('+', style='Black')
            else:
                self.paragraph.add_run(' ')
        return header

    def _add_styles(self):
        '''Sets the docx styles'''

        # create normal style
        style = self.file.styles['Normal']
        font = style.font
        font.name = 'Courier New'
        font.size = Pt(8)
        # create red style
        style = self.file.styles.add_style('Red', WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.color.rgb = RGBColor(0xFF, 0x0, 0x0)
        font.name = 'Courier New'
        font.size = Pt(8)
        # create black style
        style = self.file.styles.add_style('Black', WD_STYLE_TYPE.CHARACTER)
        font = style.font
        font.color.rgb = RGBColor(0x0, 0x0, 0x0)
        font.name = 'Courier New'
        font.size = Pt(8)

    def _get_path(self, out):
        '''Returns the path for the outfile'''

        relative = out[0] not in ['/', '~']
        if relative:
            path = os.path.join(PATH, out)
        else:
            path = out
        if self.mode == 'text' and os.path.splitext(path)[1] != '.txt':
            path = '.'.join([path, 'txt'])
        if self.mode == 'docx' and os.path.splitext(path)[1] != '.docx':
            path = '.'.join([path, 'docx'])
        return path

# ------------------
#      PROTEIN
# ------------------


class ProteinSelection(QtGui.QWidget):
    '''Protein selection widget interface'''

    def __init__(self, parent=None):
        super(ProteinSelection, self).__init__(parent)

        self.setWindowTitle("Enter Proteins")
        self.setObjectName("ProteinSelection")
        # bind instance attributes
        self.layout = QtGui.QVBoxLayout(self)
        self.layout.setAlignment(QtCore.Qt.AlignHCenter |
                                 QtCore.Qt.AlignCenter)
        header = QtGui.QLabel("Input Proteins")
        header.setAlignment(QtCore.Qt.AlignCenter)
        header.setStyleSheet(QLABEL_BANNER_STYLE)
        self.layout.addWidget(header)
        self.proteins = []
        # initialize the table and set the defaults
        self.table = Table(self)
        self.layout.addWidget(self.table)
        # add in a submit button
        submit = QtGui.QPushButton("Submit")
        submit.clicked.connect(self.parent().get_sequences)
        self.layout.addWidget(submit)
        self.show()


# ------------------
#       FILES
# ------------------


class FileSelection(QtGui.QWidget):
    '''Custom user widget with a scrollarea and files + optional conditons'''

    _entry = namedtuple("Entry", "plus file condition minus")
    _min_size = 30
    _max_size = 150
    _path = os.path.expanduser('~')

    def __init__(self, parent=None):
        super(FileSelection, self).__init__(parent)

        # init window
        self.setWindowTitle("Select Files & Conditions")
        self.setObjectName("FileSelection")
        # make layout
        self.widget_layout = QtGui.QVBoxLayout(self)
        self.scrollarea = QtGui.QScrollArea()
        self.widget = QtGui.QWidget()
        self.layout = QtGui.QVBoxLayout(self.widget)
        self.scrollarea.setWidgetResizable(True)
        self.widget_layout.addWidget(self.scrollarea)
        self.scrollarea.setWidget(self.widget)
        self.layout.setAlignment(QtCore.Qt.AlignTop | QtCore.Qt.AlignHCenter)
        # add a header
        header = QtGui.QLabel("Input Files")
        header.setAlignment(QtCore.Qt.AlignCenter)
        header.setStyleSheet(QLABEL_BANNER_STYLE)
        self.layout.addWidget(header)
        # add a horizontal widget layout
        self.hlayout = QtGui.QHBoxLayout()
        self.layout.addLayout(self.hlayout)
        self.layout.addSpacing(1)
        # make a submit button
        submit = QtGui.QPushButton("Submit")
        submit.clicked.connect(self.parent().process_output)
        self.layout.addWidget(submit)
        # make storage objs
        self.entries = []
        self.layouts = {
            'plus': QtGui.QVBoxLayout(),
            'files': QtGui.QVBoxLayout(),
            'conditions': QtGui.QVBoxLayout(),
            'minus': QtGui.QVBoxLayout()
        }
        for layout in self.layouts.values():
            layout.setAlignment(QtCore.Qt.AlignHCenter | QtCore.Qt.AlignCenter)
        self.hlayout.addLayout(self.layouts['plus'])
        self.hlayout.addLayout(self.layouts['files'])
        self.hlayout.addLayout(self.layouts['conditions'])
        self.hlayout.addLayout(self.layouts['minus'])
        # init and show
        self.make_gui()
        self.show()

    # ------------------
    #        MAIN
    # ------------------

    def make_gui(self):
        '''Makes the visual elements and inserts them into the widget'''

        # make the user widgets
        plus = QtGui.QLabel("")
        files = QtGui.QLabel("Files")
        files.setAlignment(QtCore.Qt.AlignCenter)
        files.setStyleSheet(QLABEL_STYLE)
        conditions = QtGui.QLabel("Conditions")
        conditions.setAlignment(QtCore.Qt.AlignCenter)
        conditions.setStyleSheet(QLABEL_STYLE)
        minus = QtGui.QLabel("")
        # now add all the items
        self.layouts['plus'].addWidget(plus)
        self.layouts['files'].addWidget(files)
        self.layouts['conditions'].addWidget(conditions)
        self.layouts['minus'].addWidget(minus)
        # now need to make the entries
        self.make_entry(len(self.entries))

    def make_entry(self, row):
        '''Makes a single, horizontal file entry'''

        # make user widgets
        plus = QtGui.QPushButton("+")
        plus.setMaximumWidth(self._min_size)
        plus.clicked.connect(partial(self._add_row, plus, row))
        file_btn = QtGui.QPushButton("Choose a File")
        file_btn.setMaximumWidth(self._max_size)
        conditions = QtGui.QLineEdit("Condition")
        conditions.setMaximumWidth(self._max_size)
        file_btn.clicked.connect(partial(self._get_file, file_btn, conditions))
        minus = QtGui.QPushButton("-")
        minus.setMaximumWidth(self._min_size)
        minus.clicked.connect(partial(self._delete_row, minus, row))
        # first row
        if row == 0:
            minus.setFlat(True)
        # now add all the items
        self.layouts['plus'].addWidget(plus)
        self.layouts['files'].addWidget(file_btn)
        self.layouts['conditions'].addWidget(conditions)
        self.layouts['minus'].addWidget(minus)
        # make a packaged tuple and add
        tup = self._entry(plus, file_btn, conditions, minus)
        self.entries.append(tup)

    # ------------------
    #      UTILS
    # ------------------

    def _add_row(self, widget, current_row):
        '''Adds to the row if the widget is not flat'''

        # only adds if the widget is not inactivated
        if not widget.isFlat():
            self.make_entry(len(self.entries))
        # inactivate all but end
        for row in range(0, current_row+1):
            tup = self.entries[row]
            tup.plus.setFlat(True)

    def _delete_row(self, widget, row):
        '''Deletes the row if the widget is not flat'''

        # last row, need to activate row - 1
        if row == len(self.entries) - 1:
            tup = self.entries[row-1]
            tup.plus.setFlat(False)
        if not widget.isFlat():
            tup = self.entries.pop(row)
            self.layouts['plus'].removeWidget(tup.plus)
            tup.plus.deleteLater()
            self.layouts['files'].removeWidget(tup.file)
            tup.file.deleteLater()
            self.layouts['conditions'].removeWidget(tup.condition)
            tup.condition.deleteLater()
            self.layouts['minus'].removeWidget(tup.minus)
            tup.minus.deleteLater()

    def _get_file(self, file_btn, conditions_btn):
        '''Grabs the file, and if the conditions name is unset, set it'''

        func = QtGui.QFileDialog.getOpenFileName
        dialog = func(self, 'Select file', self._path)
        file_path = dialog[0]
        if file_path:
            # store path for later
            self._path = os.path.dirname(file_path)
            name = os.path.basename(file_path)
            # set into the file
            file_btn.setText(name)
            file_btn.path = file_path
            # conditions
            if conditions_btn.text() == "Condition":
                conditions_btn.clear()
                conditions_btn.setText(name)

# ------------------
#      WIDGETS
# ------------------


class Table(QtGui.QTableWidget):
    '''Custom implementation of a QTableWidget'''

    def __init__(self, parent=None):
        super(Table, self).__init__(parent)

        self.horizontalHeader().setStyleSheet("background-color:white")
        self.horizontalHeader().setResizeMode(QtGui.QHeaderView.Stretch)
        self.verticalHeader().hide()
        self.horizontalHeader().hide()
        # need to initialize it
        self.setRowCount(1)
        self.setColumnCount(1)
        item = self._styleitem()
        self.setItem(0, 0, item)
        self.style_row(0)
        # bind signals
        self.cellChanged.connect(self.updatecells)

    # ------------------
    #       MAIN
    # ------------------

    def style_row(self, row):
        '''Sets default height for a row'''

        # set row height
        self.setRowHeight(row, ROW_HEIGHT)

    def updatecells(self, row, col):
        '''Updates the current cells and adds a row'''

        del col
        rows = self.rowCount()
        if row == rows - 1:
            # add a row
            self.setRowCount(rows + 1)
            item = self._styleitem()
            block_once(self, partial(self.setItem, rows, 0, item))
            self.style_row(rows)
        # check valid data
        # pylint: disable=bad-continuation
        item = self.item(row, 0)
        if (item is not None and
            item.text() != '' and
            len(item.text()) not in [6, 10]):
            popup = QtGui.QMessageBox(text="Warning: Please enter a valid "
                                      "uniprot ID", windowTitle='Input Error',
                                      parent=self)
            popup.exec_()

    # ------------------
    #      UTILS
    # ------------------

    @staticmethod
    def _styleitem(text=""):
        '''Creates a styled item to insert into the QTableWidget'''

        item = QtGui.QTableWidgetItem(text)
        item.setTextAlignment(QtCore.Qt.AlignHCenter |
                              QtCore.Qt.AlignCenter)
        return item

# ------------------
#  MAIN APPLICATION
# ------------------


class MainWindow(QtGui.QMainWindow):
    '''Launch Main Window'''

    sequences = None
    files = None
    dataframes = None

    def __init__(self, proteins, files, conditions=None, out=None, mode=None):
        super(MainWindow, self).__init__()

        self.proteins = proteins
        self.files = files
        self.conditions = conditions
        self.out = ARGS.output if out is None else out
        self.mode = ARGS.mode if mode is None else mode
        # init main widget
        if self.proteins is None:
            self.child_widget = ProteinSelection(self)
            self.setCentralWidget(self.child_widget)
        elif self.files is None:
            self.get_sequences()
        else:
            self.get_sequences()
            self.process_output()
        self.setStyleSheet("background-color: white")
        self.setFixedSize(400, 400)

    # ------------------
    #       MAIN
    # ------------------

    def get_sequences(self):
        '''Grabs all the UniProt sequences from a list of UniProt IDs'''

        if self.proteins is None:
            self._get_proteins()
            self.child_widget.hide()
            self.child_widget.deleteLater()
        else:
            self._check_proteins()
        if not self.proteins:
            # no proteins entered...
            self._end_error("Please enter at least one protein sequence")

        # init return
        self.sequences = []
        # grab sequences
        proteins = uniquer(self.proteins)
        for protein in proteins:
            sequence = self.connect_uniprot(protein)
            # remove header
            self.sequences.append(sequence)

        if self.files is None:
            self.child_widget = FileSelection(self)
            self.setCentralWidget(self.child_widget)

    def process_output(self):
        '''Processes and writes the output to file'''

        if self.files is None:
            self._get_files()
            self._get_dataframes()
        else:
            self._check_files()
            self._get_dataframes()
        if not self.files:
            # no files entered
            self._end_error("Please enter at least one file")

        writer = Writer(self.mode, self.out)
        if not hasattr(writer, "file"):
            self._end_error("Cannot find the save directory. Aborting...")
        for idx, sequence in enumerate(self.sequences):
            protein = self.proteins[idx]
            # print header
            sequence = sequence.splitlines()
            writer.start_sequence(sequence)
            # grab coverage conditions for each file
            coverage_list, cut_list = get_coverage(self.dataframes, protein,
                                                   sequence)
            for index in range(1, len(sequence)):
                # write sequence
                writer.write_sequence_line(sequence, index)
                # grab conditions and write coverage
                for cond_idx, condition in enumerate(self.conditions):
                    header = writer.get_header(condition)
                    # grab dataframe and map sequence coverage
                    coverage = coverage_list[cond_idx]
                    cuts = cut_list[cond_idx]
                    writer.write_condition(header, coverage, cuts,
                                           sequence, index)
                writer.write_blank_line()
        writer.close_sequence()
        writer.close()
        # now need to close the main widget
        self.close()
        sys.exit(0)

    # ------------------
    #    SEQ FUNCTIONS
    # ------------------

    @staticmethod
    def connect_uniprot(protein, host=SERVER['host']):
        '''Connects to the UniProt database and makes the query'''

        # establish connection
        if six.PY2:
            conn = httplib.HTTPConnection(host)
        else:
            conn = http.client.HTTPConnection(host)
        # make query
        url = (SERVER['path'] + SERVER['domain'] + SERVER['path'] +
               protein + SERVER['suffix'])
        # post request and read response
        conn.request('GET', url)
        response = conn.getresponse()
        value = response.read()
        if six.PY3:
            value = value.decode('utf-8')
        return value

    # ------------------
    #      UTILS
    # ------------------

    def _end_error(self, msg):
        '''Error message if an invalud sequence is entered'''

        popup = QtGui.QMessageBox(text=msg, windowTitle="Input Error",
                                  parent=self)
        popup.exec_()
        sys.exit(1)

    def _get_proteins(self):
        '''Converts the QTableWidget into a protein list'''

        self.proteins = []
        for row in range(self.child_widget.table.rowCount()):
            protein = self.child_widget.table.item(row, 0).text()
            if protein != '' and len(protein) not in [6, 10]:
                self._end_error("{0} is not a valid protein.".format(protein))
            elif protein != '':
                self.proteins.append(protein)

    def _get_files(self):
        '''Returns a list of files from a child widget'''

        self.files = []
        self.conditions = []
        for tup in self.child_widget.entries:
            if not hasattr(tup.file, "path"):
                # skip row if file never set
                continue
            self.files.append(tup.file.path)
            self.conditions.append(tup.condition.text())

    def _check_proteins(self):
        '''Ensures all the proteins are of the proper length'''

        if not all([len(i) in [6, 10] for i in ARGS.protein]):
            self._end_error('Please enter a valid UniProt ID')

    def _check_files(self):
        '''Ensures all the entered files and the conditions are right'''

        if not all([os.path.exists(i) for i in self.files]):
            self._end_error('Please enter paths to valid Protein '
                            'Prospector result files.')
        if self.conditions is not None:
            if len(self.files) != len(self.conditions):
                self._end_error('Please enter an equal number of '
                                'conditions and files.')
        else:
            self.conditions = [os.path.basename(i) for i in self.files]

    def _get_dataframes(self):
        '''Returns a list of Pandas dataframe instances from the file list'''

        try:
            self.dataframes = []
            for name in self.files:
                with open(name, 'r') as fileobj:
                    try:
                        _df = pd.read_csv(fileobj, header=2,
                                          sep='\t', engine='python')
                    except StopIteration:
                        raise AssertionError
                    # check if needed information available
                    assert PROTEIN_COLUMN in _df.columns
                    assert PEPTIDE_COLUMN in _df.columns
                    # add to dataframe list
                    self.dataframes.append(_df)
        except (IOError, OSError, AssertionError):
            basename = os.path.basename(name)
            self._end_error('{0} is not recognized. Please enter valid Protein'
                            ' Prospector result files.'.format(basename))

# ------------------
#       MAIN
# ------------------


def main():
    '''On start'''

    app = QtGui.QApplication([])
    mainwindow = MainWindow(ARGS.protein, ARGS.files, ARGS.conditions)
    mainwindow.show()
    status = app.exec_()
    sys.exit(status)

if __name__ == '__main__':
    main()
